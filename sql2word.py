import sqlite3, sys
from docx import Document


# Print iterations progress
def printProgressBar(
    iteration,
    total,
    prefix="",
    suffix="",
    decimals=1,
    length=100,
    printEnd="\r",
):
    percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
    filledLength = int(length * iteration // total)
    bar = "█" * filledLength + "-" * (length - filledLength)
    print(f"\r{prefix} |{bar}| {percent}% {suffix}", end=printEnd)
    # Print New Line on Complete
    if iteration == total:
        print()


def run_query(query, db_path):
    connect = sqlite3.connect(db_path)
    cursor = connect.cursor()
    cursor.execute(query)
    output = list([list(i) for i in cursor.fetchall()])
    cursor.execute(query)
    if cursor.description == None:
        return [[]]
    output.extend([list(map(lambda x: x[0], cursor.description))])
    connect.commit()
    connect.close()
    return output[::-1]


def run(querys, output_file_path, db_path):
    querys = [
        query[query.find("\n") :].replace("\n", " ").strip()
        for query in querys.split("-- ")[1:]
    ]
    document = Document()
    document.add_heading("SQL2Word by JJTV")
    for query, iterator in zip(querys, range(1, len(querys) + 1)):
        output = run_query(query, db_path)
        document.add_paragraph(f"Query {iterator}", style="ListBullet")
        table = document.add_table(rows=0, cols=len(output[0]), style="Table Grid")
        for row, rowiterator in zip(output, range(len(output))):
            cells = table.add_row().cells
            for column, i in zip(row, range(len(row))):
                cells[i].text = str(column)
                cells[i].paragraphs[0].runs[0].font.name = "Arial"
                if rowiterator == 0:
                    cells[i].text = str(column).upper()
                    cells[i].paragraphs[0].runs[0].font.bold = True
        printProgressBar(iterator, len(querys), "Running... ")

    document.save(output_file_path)


def main():
    args = sys.argv[1:]
    if len(args) < 3:
        print("Correct Usage\npython3 sql2word.py db.sqlite querys.sql output.docx")
        exit()
    try:
        querys = open(args[1], "r", encoding="UTF-8").read()
    except:
        print("Input File Error")
        exit()
    db_path = args[0]
    output_file_path = args[2]
    run(querys, output_file_path, db_path)


if __name__ == "__main__":
    main()
