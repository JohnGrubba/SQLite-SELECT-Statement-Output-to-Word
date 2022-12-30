# Author: Huber Elias

# pip install python-docx
# pip install db-sqlite3
import docx
from docx.shared import Pt
import sqlite3
import sys
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

erste_aufgabe = 1

# ! only SELECTs and every selected Part hat to be selected with an AS
# ! Aufgaben have to be in the right order and without holes (like 3, 4, 6, 7, ... / missing Aufgabe 5)
# ! There MUST be a "-- Aufgabe X" infront of every Query
sql_query = """-- Aufgabe 1
SELECT attribute AS "attribute" FROM table;

-- Aufgabe 2
SELECT attribute2 AS "attribute2" FROM table2;"""

db_path = "./random_path/db.sqlite"

output_file_path = "./random_path/output_file.docx"


def handle_query(query):
    connect = sqlite3.connect(db_path)
    cursor = connect.cursor()
    cursor.execute(query)
    output = cursor.fetchall()
    connect.commit()
    connect.close()
    return output


def get_list_with_querys():
    list_with_querys = sql_query.replace("\n", " ").split("-- Aufgabe ")
    list_with_querys
    del list_with_querys[0]
    for i in range(len(list_with_querys)):
        amount = 0
        for char in list_with_querys[i]:
            try:
                int(char)
                amount += 1
            except ValueError:
                break
        list_with_querys[i] = list_with_querys[i][amount:]
    return list_with_querys


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def create_header_font(table, i):
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(15)
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x7a, 0x39)


def create_row_aufgabe_x(doc, i):
    table = doc.add_table(rows=1, cols=1)

    table.rows[0].cells[0].text = "Aufgabe " + str(i)

    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(16)
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def get_attributes():
    list_with_querys = get_list_with_querys()
    attribute_list: list[list] = list()

    for query in list_with_querys:

        query = query.split("SELECT")
        query = query[1][1:]
        query = query.split("FROM")
        query = query[0][:-1]

        if "AS" not in query:
            print("go fuck yourself")
            print("every selected attribute has to be selected with an AS")
            sys.exit()
        else:
            query = query.split('AS "')
            for part in range(len(query)):
                query[part] = query[part].split(", ")
            attributes = list()
            new_query = list()
            for part in query:
                for word in part:
                    new_query.append(word)
            query = new_query
            for part in range(len(query)):
                if part % 2 != 0:
                    attributes.append(query[part].replace('"', ""))
            attribute_list.append(attributes)

    return attribute_list


def create_new_table(query_output, doc, i_gesamt):
    table = doc.add_table(rows=len(query_output)+1, cols=len(query_output[0]))

    # Header Row
    i_gesamt = i_gesamt - erste_aufgabe
    attributes = get_attributes()

    for i in range(len(query_output[0])):
        table.rows[0].cells[i].text = attributes[i_gesamt][i]
        create_header_font(table, i)

    # Entitys
    entity_i = 0
    data_i = 0
    for entity in query_output:
        for data in entity:
            table.rows[entity_i+1].cells[data_i].text = str(data)

            data_i += 1
        data_i = 0
        entity_i += 1


def main():
    # Open Word File
    try:
        doc = docx.Document(output_file_path)
    except docx.opc.exceptions.PackageNotFoundError:
        doc = docx.Document()

    # Clear Word File
    # Tables
    for activeTable in doc.tables:
        activeTable._element.getparent().remove(activeTable._element)
    # Patagraphs
    for para in doc.paragraphs:
        delete_paragraph(para)

    # Get List with the Querys
    list_with_querys = get_list_with_querys()

    i = erste_aufgabe
    for query in list_with_querys:
        # Execute Query
        output = handle_query(query)

        # Save to Word File
        create_row_aufgabe_x(doc, i)
        create_new_table(output, doc, i)
        i += 1

    # Save Word File
    doc.save(output_file_path)
    return 0


if __name__ == "__main__":
    main()
